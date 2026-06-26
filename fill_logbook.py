import docx
from docx.shared import Pt, Inches

def create_filled_logbook():
    doc = docx.Document('Log book.docx')
    
    # We will assume the first table is a template. We'll extract its style if needed, 
    # but python-docx has a hard time copying tables perfectly.
    # Instead, we will add new tables for each entry.
    
    entries = [
        {
            "date": "January 2, 2026",
            "carried": "- Conducted initial literature review and researched topics related to a Blockchain-Based Community-Assisted Privacy-Preserving Reporting Service for Local Governance.\n- Explored methodologies to address user privacy, anonymity, and data integrity in e-governance systems using cryptographic approaches.\n- Investigated potential solutions for handling user trustworthiness and reviewed existing crowdsourcing frameworks.",
            "planned": "- Define the overall system architecture and formulate detailed workflows for the core phases (User Registration, Reporting, and Opinion Polls).\n- Determine the appropriate technology stack (front-end frameworks, blockchain platform, etc.) needed for the implementation.\n- Draft the initial sections of the project proposal."
        },
        {
            "date": "January 20, 2026",
            "carried": "- Finalized system architectures and operational workflows for User Registration (using LES), Reporting, and Opinion Polls.\n- Selected the technology stack, opting for Next.js for the web DApp, a permissioned blockchain for reporting, and IPFS for decentralized storage.\n- Completed the drafting of the formal project proposal incorporating the designed architectures and proposed technologies.",
            "planned": "- Finalize and submit the formal project proposal.\n- Initialize the version control system (Git) and set up the foundational project repository.\n- Establish the base directory structures for the web DApp implementation."
        },
        {
            "date": "January 29, 2026",
            "carried": "- Successfully finalized and submitted the project proposal.\n- Set up the Git version control repository and established the base project structures for the web DApp.\n- Configured initial development environments and dependencies required for front-end development and Web3 integration.",
            "planned": "- Begin the UI/UX design phase for the web DApp.\n- Start implementation of the decentralized application, focusing on the foundational layouts and components.\n- Plan and initiate the implementation of the local wallet creation process for citizen nodes."
        },
        {
            "date": "February 19, 2026",
            "carried": "- Commenced UI/UX design and implemented foundational UI components for the web DApp.\n- Implemented core cryptographic features, including local wallet creation, cryptographic hashing mechanisms, and secure digital message signings.\n- Developed and integrated the logic for connecting the DApp to the permissioned blockchain and successfully retrieving reports/tickets to display on the interface.",
            "planned": "- Develop the functionality for sending new reports with data payloads to the blockchain.\n- Integrate Zero-Knowledge Proof (ZKP) based GovID authentication for privacy-preserving user logins.\n- Refine and test the report submission workflows and enhance DApp UI responsiveness."
        },
        {
            "date": "April 29, 2026",
            "carried": "- Finalized the report submission workflows, enabling the sending of reports with complex data payloads to the blockchain.\n- Successfully integrated the ZKP-based GovID authentication page to ensure privacy-preserving citizen login and verification.\n- Refined DApp UI layouts with responsive navigation and completed UI changes for better user experience.\n- Conducted comprehensive testing on report submission, retrieval workflows, and wallet interactions.",
            "planned": "- Finalize the project progress documentation and prepare the final presentation.\n- Conduct complete end-to-end integration testing across the web DApp, AI oracles, relayer, and smart contracts.\n- Address any minor UI bugs and optimize DApp performance."
        }
    ]
    
    # Let's populate the first table with the first entry
    if len(doc.tables) > 0:
        table = doc.tables[0]
        table.cell(0, 1).text = entries[0]["date"]
        table.cell(3, 0).text = entries[0]["carried"]
        table.cell(5, 0).text = entries[0]["planned"]
        
    # For subsequent entries, we'll try to add new tables that look similar
    for entry in entries[1:]:
        doc.add_paragraph() # spacing
        new_table = doc.add_table(rows=7, cols=6)
        new_table.style = 'Table Grid'
        
        # Row 0
        new_table.cell(0, 0).text = 'Meeting Date:'
        new_table.cell(0, 1).text = entry['date']
        new_table.cell(0, 2).text = 'Semester:'
        new_table.cell(0, 4).text = 'Week:'
        
        # Row 2 (headers)
        hdr1 = new_table.cell(2, 0)
        hdr1.merge(new_table.cell(2, 5))
        hdr1.text = 'Work carried out since the last meeting'
        
        # Row 3 (content)
        cnt1 = new_table.cell(3, 0)
        cnt1.merge(new_table.cell(3, 5))
        cnt1.text = entry['carried']
        
        # Row 4 (headers)
        hdr2 = new_table.cell(4, 0)
        hdr2.merge(new_table.cell(4, 5))
        hdr2.text = 'Work planned for the coming week'
        
        # Row 5 (content)
        cnt2 = new_table.cell(5, 0)
        cnt2.merge(new_table.cell(5, 5))
        cnt2.text = entry['planned']
        
        # Row 6 (footer)
        ftr = new_table.cell(6, 0)
        ftr.merge(new_table.cell(6, 5))
        ftr.text = 'Signature of the supervisor:'
        
    doc.save('Filled_Log_book.docx')
    print("Saved as Filled_Log_book.docx")

if __name__ == '__main__':
    create_filled_logbook()
